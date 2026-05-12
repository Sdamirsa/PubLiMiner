#!/usr/bin/env python3
"""Generate a Word document report for a PubLiMiner per-project analysis.

Reads output/<project>/analysis/figure_captions.json and the PNG figures,
then produces a formatted .docx with:
  - Page 1: project title + 2-paragraph introduction
  - One page per figure: image (full-width), bold title, visualization + insight
  - Last page: condensed take-home paragraph

Usage:
    python scripts/per-project/generate_word_report.py --project-dir output/cardiac_ct
    python scripts/per-project/generate_word_report.py --project-dir output/cardiac_mri

Output: output/<project>/analysis/report_<project_name>.docx
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ── Figure display order ───────────────────────────────────────────────────────
FIGURE_ORDER = [
    "T1_volume_by_year",
    "T2_first_author_specialty_by_year",
    "T3_last_author_specialty_by_year",
    "T4_corresponding_author_specialty_by_year",
    "T5_study_type_by_year",
    "T6_majority_specialty",
    "G1_country_map",
    "G1b_country_map_world",
    "G1c_country_map_APC",
    "G2_top20_countries",
    "G3_region_over_time",
    "G4_decade_comparison",
    "I1_top20_institutions_global",
    "I2_top20_institutions_european",
    "I3_institutions_by_role",
    "I4_institution_by_country",
    "J1_top20_journals_all",
    "J2_top20_journals_european",
    "X1_specialty_studytype_heatmap",
    "X2_dashboard_overview",
    "bern_stat",
]


# ── Project metadata ───────────────────────────────────────────────────────────
_META: dict[str, dict] = {
    "cardiac_ct": {
        "doc_title": "Cardiac CT — Global Publication Analysis 2001–2026",
        "intro": [
            (
                "This report summarises a systematic bibliometric analysis of the global cardiac "
                "computed tomography (cardiac CT) literature indexed on PubMed between 2001 and "
                "2026. The primary aim was to characterise the scope, growth trajectory, "
                "geographic distribution, institutional landscape, and disciplinary composition "
                "of this field, with particular emphasis on European and Swiss contributions."
            ),
            (
                "All PubMed records matching a cardiac CT search query were retrieved and "
                "deduplicated (n = 16,449 unique papers), then processed through an automated "
                "pipeline: XML parsing of author affiliations for country, institution, and "
                "specialty inference; large-language-model (LLM) structured extraction of study "
                "type and first/last/corresponding-author specialty; and a geographic filter to "
                "identify papers with at least one European-affiliated author (n = 6,312; 38.4%). "
                "All figures are generated on the full extracted corpus without relevance "
                "filtering. No manual curation of the paper set was performed. Minor "
                "over-attribution may be present due to multi-national authorships."
            ),
        ],
        "conclusion": (
            "The cardiac CT literature is large, rapidly expanding, geographically concentrated, "
            "and predominantly descriptive in study design. From near zero in 2001, annual output "
            "reached approximately 1,600 papers by 2024, driven by the clinical adoption of "
            "coronary CT angiography from ~2012. Cardiology holds ~65% of first-author "
            "designations with Radiology stable at ~25–30%, confirming a dual-specialty field "
            "tilted toward Cardiology. Geographically, the USA leads in absolute volume, but East "
            "Asian nations — Japan, China, South Korea — have collectively risen to match or "
            "exceed North American output by the late 2010s. Case reports and series remain the "
            "dominant study design (~40–50%), with original research at only ~10–13%, a ratio "
            "that has not improved over 25 years. A small number of centres accounts for "
            "disproportionate output: Leiden UMC leads European hospitals, and within Switzerland, "
            "Inselspital ranks 9th (65 papers) and the University of Bern 13th (53 papers) among "
            "European institutions — a strong position for a mid-sized academic health system."
        ),
    },
    "cardiac_mri": {
        "doc_title": "Cardiac MRI — Global Publication Analysis 2001–2026",
        "intro": [
            (
                "This report presents a systematic bibliometric analysis of the global cardiac "
                "magnetic resonance imaging (cardiac MRI / CMR) literature indexed on PubMed "
                "between 2001 and 2026. The primary aim was to characterise the field's growth "
                "trajectory, geographic and institutional distribution, disciplinary composition, "
                "and publication patterns, with particular emphasis on European and Swiss "
                "contributions."
            ),
            (
                "All PubMed records matching a cardiac MRI search query were retrieved and "
                "deduplicated (n = 35,512 unique papers), then processed through the same "
                "automated pipeline used for the cardiac CT analysis: XML parsing of author "
                "affiliations for country, institution, and specialty inference; LLM extraction of "
                "study type and author specialty; and a geographic filter to identify "
                "European-affiliated papers (n = 16,823; 47.4%). All figures are generated on the "
                "full extracted corpus without relevance filtering. No manual curation was applied; "
                "minor over-attribution from multi-national authorships may be present."
            ),
        ],
        "conclusion": (
            "The cardiac MRI literature is substantially larger than cardiac CT (34,844 vs 16,449 "
            "papers), reflecting the modality's broader application across anatomy, function, tissue "
            "characterisation, and congenital disease. Annual output grew from under 100 papers in "
            "2001 to approximately 2,700-3,000 per year by 2022-2025. Cardiology holds 77% of "
            "first-author designations -- more dominant than in cardiac CT (64%) -- while "
            "Radiology's share (~15%) is comparable across both modalities, confirming that cardiac "
            "MRI is firmly embedded in cardiology academic culture. The USA leads at 10,244 papers, "
            "but Germany (3,366) and Italy (2,729) are prominent second and fourth -- a more "
            "European-centred geographic distribution than cardiac CT, consistent with the higher "
            "European share (47.4% vs 38.4%). Case reports and series remain the dominant study "
            "design (~43%), with original research at ~11%, virtually identical to the cardiac CT "
            "pattern despite a much larger corpus. The Journal of Cardiovascular Magnetic Resonance "
            "leads at 1,778 papers (~5% of the entire corpus), serving as the field's dedicated "
            "flagship. Within Switzerland, Inselspital ranks 8th among 4,949 European hospitals "
            "(136 papers) and the University of Bern ranks 7th among 2,383 European universities "
            "(110 papers, 154 combined) -- a stronger European position than in cardiac CT, "
            "confirming Bern as one of the most active European academic centres in both cardiac "
            "imaging modalities."
        ),
    },
}


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_margins(doc: Document, cm_val: float = 1.5) -> None:
    """Set all page margins uniformly."""
    for section in doc.sections:
        margin = Cm(cm_val)
        section.top_margin    = margin
        section.bottom_margin = margin
        section.left_margin   = margin
        section.right_margin  = margin


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def _heading(doc: Document, text: str, level: int = 1,
             color: RGBColor | None = None) -> None:
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def _body(doc: Document, text: str, italic: bool = False,
          size_pt: int = 10, space_after: int = 6) -> None:
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.font.size   = Pt(size_pt)
    run.font.italic = italic
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(space_after)


def _add_figure_page(
    doc: Document,
    img_path: Path,
    caption: dict,
) -> None:
    """Add one page: figure image, bold title, visualization (italic), insight."""
    # Figure image — constrain to page width (A4 portrait with 1.5cm margins ≈ 17.7cm = 6.97")
    max_width = Inches(6.9)
    try:
        doc.add_picture(str(img_path), width=max_width)
        # centre the image
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_para.paragraph_format.space_before = Pt(0)
        last_para.paragraph_format.space_after  = Pt(6)
    except Exception as exc:
        _body(doc, f"[Image not found: {img_path.name}]")

    # Title (bold, 12pt)
    p = doc.add_paragraph()
    run = p.add_run(caption.get("title", img_path.stem))
    run.font.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)

    # Visualization description (italic, 9.5pt)
    viz = caption.get("visualization", "")
    if viz:
        _body(doc, viz, italic=True, size_pt=9, space_after=3)

    # Insight paragraph (normal, 9.5pt)
    ins = caption.get("insight", "")
    if ins:
        _body(doc, ins, italic=False, size_pt=9, space_after=4)


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    import json

    parser = argparse.ArgumentParser(description="Generate Word report from figure captions")
    parser.add_argument("--project-dir", required=True,
                        help="Project output directory (e.g. output/cardiac_ct)")
    args = parser.parse_args()

    base_dir   = Path(args.project_dir).resolve()
    analysis   = base_dir / "analysis"
    fig_dir    = analysis / "figures"
    caps_path  = analysis / "figure_captions.json"

    if not caps_path.exists():
        raise SystemExit(f"figure_captions.json not found at {caps_path}")

    with caps_path.open(encoding="utf-8") as f:
        captions: dict[str, dict] = json.load(f)

    # Derive project key from directory name
    proj_key  = base_dir.name   # "cardiac_ct" or "cardiac_mri"
    proj_name = " ".join(
        w.upper() if w.lower() in ("ct", "mri", "pet") else w.title()
        for w in proj_key.replace("_", " ").split()
    )
    meta = _META.get(proj_key, {
        "doc_title": f"{proj_name} — Publication Analysis",
        "intro": [
            f"Bibliometric analysis of {proj_name} publications on PubMed (2001–2026).",
        ],
        "conclusion": None,
    })

    # Build conclusion from captions JSON if not hardcoded
    conclusion = meta.get("conclusion")
    if conclusion is None:
        conclusion = captions.get("_conclusion", {}).get("text") or (
            f"See individual figure captions for detailed findings from the {proj_name} analysis."
        )

    doc = Document()
    _set_margins(doc, cm_val=1.5)

    # ── Cover / intro page ────────────────────────────────────────────────────
    _heading(doc, meta["doc_title"], level=1,
             color=RGBColor(0x1F, 0x4E, 0x79))

    intro_paras = meta.get("intro", [])
    if isinstance(intro_paras, str):
        intro_paras = [intro_paras]
    for para_text in intro_paras:
        _body(doc, para_text, size_pt=11, space_after=8)

    _add_page_break(doc)

    # ── Figure pages ──────────────────────────────────────────────────────────
    # Ordered figures first, then any remaining in JSON order
    ordered_keys = [k for k in FIGURE_ORDER if k in captions]
    extra_keys   = [k for k in captions if k not in FIGURE_ORDER and not k.startswith("_")]
    all_keys     = ordered_keys + extra_keys

    for i, key in enumerate(all_keys):
        img_path = fig_dir / f"{key}.png"
        caption  = captions[key]
        _add_figure_page(doc, img_path, caption)
        if i < len(all_keys) - 1:
            _add_page_break(doc)

    # ── Conclusion page ───────────────────────────────────────────────────────
    _add_page_break(doc)
    _heading(doc, "Key Takeaways", level=1,
             color=RGBColor(0x1F, 0x4E, 0x79))
    _body(doc, conclusion, size_pt=11, space_after=8)

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = analysis / f"report_{proj_key}.docx"
    doc.save(str(out_path))
    print(f"[OK] Saved -> {out_path}  ({out_path.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    main()
